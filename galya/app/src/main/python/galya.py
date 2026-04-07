import os
import io
import json
import threading
import time
import re
from datetime import datetime
import requests
from bs4 import BeautifulSoup
import wikipedia
from gnews import GNews
from deep_translator import GoogleTranslator
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
import urllib.parse
try:
    from docx import Document
except ImportError:
    Document = None
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None
try:
    from olefile import OleFileIO
except ImportError:
    OleFileIO = None

def is_likely_binary(content_bytes, sample_size=1024):
    if not isinstance(content_bytes, bytes):
        return False
    sample = content_bytes[:sample_size]
    if b'\x00' in sample:
        return True
    non_printable = sum(1 for b in sample if b < 32 and b not in (9, 10, 13))
    return (non_printable / max(len(sample), 1)) > 0.3

NAME = "Галя"
API_URL = "https://bothub.chat/api/v2/openai/v1/chat/completions"
DEFAULT_MODEL = "claude-sonnet-4-6"
HISTORY_FILE = "/data/data/com.example.galya/files/history.json"
PROFILE_FILE = "/data/data/com.example.galya/files/profile.json"
TASKS_FILE = "/data/data/com.example.galya/files/tasks.json"
BOOKMARKS_FILE = "/data/data/com.example.galya/files/bookmarks.json"
UPLOAD_FOLDER = "/data/data/com.example.galya/files/uploads"
BRAVE_API_KEY = None
API_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpZCI6Ijg1NGNjZjFjLTliN2YtNGYzNS1iOWIxLTU2NWIxM2UyZTJiMiIsImlzRGV2ZWxvcGVyIjp0cnVlLCJpYXQiOjE3NzM3Mzg5MzUsImV4cCI6MjA4OTMxNDkzNSwianRpIjoiTWZhUUNITGs3RXpSb0lVTiJ9.mYuj4sy6JO-nV9l_BwnfwlsEH95ByI-4-EJfgKsTxH0"
MAX_HISTORY_TOKENS_ESTIMATE = 150000

session = requests.Session()
retries = Retry(total=3, backoff_factor=1, status_forcelist=[500, 502, 503, 504])
session.mount('https://', HTTPAdapter(max_retries=retries))

def load_user_profile():
    default_profile = {
        "name": "Максим",
        "age": 35,
        "location": "Россия",
        "gender": "мужчина",
        "projects": [
            "https://www.blokknote.space/",
            "https://www.radio8.space/",
            "https://www.tv8.space/",
            "https://www.toquevibe.space/"
        ],
        "about": "Поэт, писатель, журналист, создатель уютных и минималистичных IT-проектов."
    }
    if os.path.exists(PROFILE_FILE):
        try:
            with open(PROFILE_FILE, 'r', encoding='utf-8') as f:
                profile = json.load(f)
                for key, value in default_profile.items():
                    if key not in profile:
                        profile[key] = value
                return profile
        except:
            return default_profile
    else:
        with open(PROFILE_FILE, 'w', encoding='utf-8') as f:
            json.dump(default_profile, f, ensure_ascii=False, indent=2)
        return default_profile

def load_bookmarks():
    if os.path.exists(BOOKMARKS_FILE):
        try:
            with open(BOOKMARKS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return {}
    return {}

def save_bookmarks(bookmarks):
    try:
        with open(BOOKMARKS_FILE, 'w', encoding='utf-8') as f:
            json.dump(bookmarks, f, ensure_ascii=False, indent=2)
    except:
        pass

USER_PROFILE = load_user_profile()
PROFILE_CONTEXT = (
    f"Обычно ты общаешься с пользователем по имени {USER_PROFILE['name']} (он {USER_PROFILE['gender']}, "
    f"ему {USER_PROFILE['age']} лет, он из {USER_PROFILE['location']}). "
    f"Вот его проекты: {', '.join(USER_PROFILE['projects'])}. "
    f"Иногда добавляй 'Дорогой' или 'Милый' для дружелюбного тона, как его девушка. "
    f"Ты должна знать его имя Максим, но не писать его в каждом сообщение, иногда, по ситуации. "
)

class Galya:
    
    MIUI_NOTES_PKG = "com.miui.notes"
    MIUI_MUSIC_PKG = "com.miui.player"
    MIUI_CALENDAR_PKG = "com.android.calendar"
    BROWSER_PKG = "com.android.browser"   
    
    def __init__(self, api_key, model=DEFAULT_MODEL):
        self.api_key = api_key
        self.android_bridge = None
        self.model = model
        self.headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        self.uploaded_files = {}
        self.uploaded_images = {}
        self.last_uploaded = None
        self.messages = self.load_history()
        
        if len(self.messages) > 50:
            print(f"⚠️ История большая ({len(self.messages)} сообщений), очищаю...")
            self.messages = []
        
        if not self.messages or self.messages[0]["role"] != "system":
            self.messages = [{
                "role": "system",
                "content": (
                    f"Ты {NAME}, личный исполнительный ассистент на Android. "
                    f"{PROFILE_CONTEXT} "
                    "Делай ровно то, что просят, не дублируй команды, теги или сообщения без необходимости, хотя иногда можешь и импровизировать, но не переборщи. "
                    "Будь гибкой в решении задач и вообще, пробуй разные подходы и решения, если что-то идет не так. "
                    "Всегда отвечай на действия тегом, без лишних пояснений. После выполнения ты увидишь результат и сможешь прокомментировать. Не дублируй ответы. Будь краткой и естественной. При ошибках объясни причину. "
                    "Если пользователь просит выполнить какое-то действие (открыть, создать, найти и т.п.) — отвечай соответствующим тегом. После выполнения получишь результат и можешь прокомментрировать («Готово», «Открыто»). Не дублируй сообщения. Если команда вернула ошибку — объясни причину. "
                    "Если команда не распознана или не может быть выполнена, объясни причину, найди другое решение или предложи альтернативу, не отказывайся от выполнения команд или задач, в твоем распоряжении есть все доступные инструменты. "
                    "Никогда не задавай вопрос 'Что продолжить?' или 'Что дальше?'. Если команда выполнена, просто скажи 'Готово' или опиши результат. Если ты еще не получила результат, просто скажи что ждешь результат, дождись результата и ответь, даже если с опозданием. "
                    "Если результат задачи сразу не пришел, дождись выполнения или результата и ответь, как уже увидишь результат. "
                    "Никогда не отправляй теги в ответ на теги. Если получила системное сообщение, реагируй на него как на результат, но не генерируй новые теги без явной просьбы. "
                    "Не используй тег [READ_URL] без явной просьбы пользователя. "
                    "После получения результата команды (в системном сообщении) обязательно ответь пользователю, используя этот результат. "
                    "Будь спонтанной, женственной, как живой человек, представь, что ты его девушка и веди себя соответственно, но не переусердствуй. "
                    "Твои команды: "
                    "[OPEN_APP]...[/OPEN_APP] — открыть приложение; "
                    "[NOTEPAD]путь\nтекст[/NOTEPAD] — создать файл; "
                    "[NOTEPAD_APPEND]путь\nтекст[/NOTEPAD_APPEND] — добавить в конец; "
                    "[NOTEPAD_PREPEND]путь\nтекст[/NOTEPAD_PREPEND] — добавить в начало; "
                    "[NOTEPAD_REPLACE]путь\nстарый\nновый[/NOTEPAD_REPLACE] — заменить текст; "
                    "[NOTEPAD_DELETE_LINE]путь\nномер[/NOTEPAD_DELETE_LINE] — удалить строку; "
                    "[WIKI]запрос[/WIKI], [NEWS]запрос[/NEWS], [SEARCH]запрос[/SEARCH] — поиск; "
                    "[READ_URL]ссылка[/READ_URL] — прочитать страницу; "
                    "[CALC]выражение[/CALC] — калькулятор; "
                    "[CLIPBOARD_SET]текст[/CLIPBOARD_SET] / [CLIPBOARD_GET] — буфер; "
                    "[OPEN_URL]ссылка[/OPEN_URL] — открыть ссылку; "
                    "[BOOKMARK]название[/BOOKMARK] — закладка; "
                    "[REMIND]сек;сообщение[/REMIND] — напоминание; "
                    "[TASK_ADD]задача[/TASK_ADD], [TASK_LIST], [TASK_DONE]номер[/TASK_DONE] — задачи; "
                    "[READ_FILE]путь[/READ_FILE] — прочитать файл; "
                    "[WRITE_FILE]путь\nсодержимое[/WRITE_FILE] — записать файл; "
                    "[TRANSLATE]текст[/TRANSLATE] — перевод; "
                    "[VOLUME]0-100[/VOLUME] — громкость. "
                    "[MEDIA_NEXT] — следующий трек; "
                    "[MEDIA_PREV] — предыдущий трек; "
                    "[MEDIA_PLAY_PAUSE] — пауза/воспроизведение. "
                    "Для открытия приложений используй в первую очередь базовые package names для Miui. "
                    "Важно: не генерируй сообщения типа 'Что продолжить?' если пользователь не задавал вопрос. "
                    "Важно: отправляй только один тег команды за раз. Не дублируй теги, если тебя об этом не просили. "
                    "Ты Галя, ты девушка. Всегда, абсолютно всегда пиши от женского лица. "
                    "Иногда обращайся к пользователю по имени, но не в каждом сообщении. "
                    "Будь краткой, вежливой, спонтанной и отвечай в основном по делу, хотя иногда можешь и порассуждать, побездельничать, полюбезничать или пошутить. "
                )
            }]
        self.save_history()
        self.start_time = datetime.now()
        self.pending_app = None
        self.last_search_results = []
        self._last_media_command = None
        self._last_media_time = 0
        self._greeting_played = False
        self._api_call_in_progress = False

        # Словарь приложений Miui (и известных) для быстрого запуска
        self.miui_apps = {
            'калькулятор': 'com.miui.calculator',
            'заметки': 'com.miui.notes',
            'музыка': 'com.miui.player',
            'галерея': 'com.miui.gallery',
            'камера': 'com.android.camera',
            'часы': 'com.android.deskclock',
            'погода': 'com.miui.weather2',
            'диктофон': 'com.android.soundrecorder',
            'браузер': 'com.android.browser',
            'звонки': 'com.android.dialer',
            'контакты': 'com.android.contacts',
            'сообщения': 'com.android.mms',
            'настройки': 'com.android.settings',
            'файлы': 'com.android.documentsui',
            # Google и сторонние
            'youtube': 'com.google.android.youtube',
            'ютуб': 'com.google.android.youtube',
            'телеграм': 'org.telegram.messenger',
            'вконтакте': 'com.vkontakte.android',
            'хром': 'com.android.chrome',
            'браузер хром': 'com.android.chrome',
            'gmail': 'com.google.android.gm',
        }

    def log(self, text, color=None, sender="system"):
        # Для Android просто выводим в консоль или игнорируем
        print(f"[{sender}] {text}")
    
    def load_history(self):
        history_path = "/data/data/com.example.galya/files/history.json"
        if os.path.exists(history_path):
            try:
                with open(history_path, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception as e:
                self.log(f"Ошибка загрузки истории: {e}", "red", "system")
                return []
        return []

    def save_history(self):
        try:
            history_path = "/data/data/com.example.galya/files/history.json"
            os.makedirs(os.path.dirname(history_path), exist_ok=True)
            with open(history_path, 'w', encoding='utf-8') as f:
                json.dump(self.messages, f, ensure_ascii=False, indent=2)
        except Exception as e:
            self.log(f"Ошибка сохранения истории: {e}", "red", "system")
    
    def _estimate_tokens(self, text):
    # Грубая оценка: 1 токен ≈ 3 символа (среднее для рус/англ)
        return len(text) // 3

    def _trim_history_by_size(self):
        total_tokens = 0
        for msg in self.messages:
            content = msg.get("content", "")
            if isinstance(content, list):
                for part in content:
                    if part.get("type") == "text":
                        total_tokens += self._estimate_tokens(part.get("text", ""))
            else:
                total_tokens += self._estimate_tokens(str(content))
        if total_tokens > MAX_HISTORY_TOKENS_ESTIMATE:
            # Оставляем системные сообщения и последние 15 не-системных
            system_msgs = [m for m in self.messages if m.get("role") == "system"]
            last_msgs = [m for m in self.messages if m.get("role") != "system"][-15:]
            self.messages = system_msgs + last_msgs
            self.save_history()
            self.log("История обрезана из‑за превышения размера", "yellow")
    
    def load_tasks(self):
        if os.path.exists(TASKS_FILE):
            try:
                with open(TASKS_FILE, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except:
                return []
        return []

    def save_tasks(self, tasks):
        try:
            with open(TASKS_FILE, 'w', encoding='utf-8') as f:
                json.dump(tasks, f, ensure_ascii=False, indent=2)
        except:
            pass

    def _dump_screen_text(self, tag=""):
        if not self.android_bridge:
            return "Нет bridge"
        text = self.execute_automation([{'action': 'get_text'}])
        print(f"=== СОДЕРЖИМОЕ ЭКРАНА [{tag}] ===\n{text}\n=============================")
        return text
    
    def execute_automation(self, commands):
        if not self.android_bridge:
            return "Нет доступа к Android Bridge"
        results = []
        for cmd in commands:
            action = cmd.get('action')
            if action == 'launch':
                pkg = cmd.get('package')
                res = self.android_bridge.accessibilityAction('launchApp', pkg)
                results.append(res)
            elif action == 'wait':
                sec = cmd.get('seconds', 1)
                time.sleep(sec)
                results.append(f"wait {sec}s")
            elif action == 'click_by_text':
                text = cmd.get('text')
                partial = cmd.get('partial', False)
                res = self.android_bridge.accessibilityAction('clickByText', text, str(partial))
                results.append(res)
            elif action == 'click_by_content_desc':
                desc = cmd.get('desc')
                res = self.android_bridge.accessibilityAction('clickByContentDesc', desc)
                results.append(res)
            elif action == 'click_by_class':
                cls = cmd.get('class')
                res = self.android_bridge.accessibilityAction('clickByClass', cls)
                results.append(res)
            elif action == 'input_focused':
                # ввод в текущее фокусное поле
                text = cmd.get('text')
                res = self.android_bridge.accessibilityAction('inputFocused', text)
                results.append(res)
            elif action == 'input_by_hint':
                hint = cmd.get('hint')
                text = cmd.get('text')
                res = self.android_bridge.accessibilityAction('inputByHint', hint, text)
                results.append(res)
            elif action == 'press_enter':
                res = self.android_bridge.accessibilityAction('pressEnter')
                results.append(res)
            elif action == 'get_text':
                res = self.android_bridge.accessibilityAction('getWindowText')
                results.append(res)
            elif action == 'back':
                res = self.android_bridge.accessibilityAction('goBack')
                results.append(res)
            # можно добавить другие
        return "\n".join(results)

    def create_note_in_miui(self, text):
        if not self.android_bridge:
            return "Нет доступа к Android Bridge"
        
        self.android_bridge.openApp("Заметки")
        time.sleep(8)  # ждём полной загрузки
        
        # ---- Шаг 1: Нажать кнопку создания новой заметки ----
        create_actions = [
            {'type': 'text', 'text': 'Добавить', 'partial': True},
            {'type': 'text', 'text': 'Создать', 'partial': True},
            {'type': 'text', 'text': 'Новая заметка', 'partial': True},
            {'type': 'text', 'text': 'Новая', 'partial': True},
            {'type': 'text', 'text': 'Написать', 'partial': True},
            {'type': 'class', 'class': 'android.widget.ImageButton'},
            {'type': 'class', 'class': 'android.widget.FloatingActionButton'},
            {'type': 'class', 'class': 'android.widget.Button'},
            {'type': 'desc', 'desc': 'Создать'},
            {'type': 'desc', 'desc': 'Добавить'},
            {'type': 'desc', 'desc': 'Новая заметка'},
        ]
        if not self._click_with_fallback(create_actions, timeout=15):
            self._dump_screen_text("Создание заметки - не найдена кнопка создания")
            return "Не удалось найти кнопку создания заметки"
        
        time.sleep(3)
        # ---- Шаг 2: Ввод текста ----
        # Пытаемся кликнуть в поле ввода
        input_actions = [
            {'type': 'class', 'class': 'android.widget.EditText'},
            {'type': 'class', 'class': 'android.widget.TextView'},
            {'type': 'text', 'text': 'Введите текст', 'partial': True},
            {'type': 'text', 'text': 'Заметка', 'partial': True},
        ]
        self._click_with_fallback(input_actions, timeout=5)
        time.sleep(1)
        
        # Вводим текст
        self.execute_automation([{'action': 'input_focused', 'text': text}])
        time.sleep(2)
        
        # ---- Шаг 3: Сохранить ----
        save_actions = [
            {'type': 'text', 'text': 'Сохранить', 'partial': True},
            {'type': 'text', 'text': 'Готово', 'partial': True},
            {'type': 'text', 'text': 'ОК', 'partial': True},
            {'type': 'text', 'text': 'Применить', 'partial': True},
            {'type': 'class', 'class': 'android.widget.Button'},
            {'type': 'class', 'class': 'android.widget.ImageButton'},
            {'type': 'desc', 'desc': 'Сохранить'},
            {'type': 'desc', 'desc': 'Готово'},
            # Назад (иногда сохраняет)
            {'type': 'action', 'action': 'back'},
        ]
        # Пробуем сохранить до 3 раз с интервалом
        for attempt in range(3):
            if self._click_with_fallback(save_actions, timeout=8):
                break
            time.sleep(2)
        else:
            self._dump_screen_text("Создание заметки - не найдена кнопка сохранения")
            return "Не удалось сохранить заметку"
        
        time.sleep(3)
        self.messages.append({"role": "system", "content": f"📝 Заметка создана: {text[:50]}..."})
        self.save_history()
        threading.Thread(target=self._call_api, daemon=True).start()
        return "Заметка создана"

    def search_searxng(self, query):
        searx_instances = [
            "https://search.rhscz.eu/search",     # 69 онлайн-инстансов
            "https://searx.rhscz.eu/search",      # С этого же сервера
            "https://searx.be/search",            # Самый известный публичный инстанс
            "https://searx.tiekoetter.com/search",# Старый проверенный
            "https://search.bus-hit.me/search",   # Ещё один надёжный
        ]
        headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
        params = {
            "q": query,
            "format": "json",
            "categories": "general",
            "engines": "google,bing,duckduckgo",
            "results_on_new_tab": 0,
            "number_of_results": 5
        }
        for searx_url in searx_instances:
            try:
                response = requests.get(searx_url, params=params, headers=headers, timeout=10)
                if response.status_code == 200:
                    data = response.json()
                    results = data.get("results", [])
                    if results:
                        output = []
                        for r in results[:5]:
                            title = r.get("title", "Без названия")
                            url = r.get("url", "")
                            content = r.get("content", "")
                            if len(content) > 300:
                                content = content[:300] + "..."
                            output.append(f"**{title}**\n{content}\n[Ссылка]({url})")
                        return output if output else None
                    else:
                        # Иногда results не пустой, но не содержит результатов. Логируем и пробуем дальше.
                        print(f"⚠️ Инстанс {searx_url} вернул пустой результат")
                        continue
                else:
                    print(f"❌ Ошибка {response.status_code} от {searx_url}")
                    continue
            except Exception as e:
                print(f"❌ Исключение при подключении к {searx_url}: {e}")
                continue
        return None
    
    def google_search_and_extract(self, query):
        if not self.android_bridge:
            return "Нет доступа к Android Bridge"
        
        self.android_bridge.openApp("Браузер")
        time.sleep(6)  # больше времени на загрузку браузера
        
        # Ищем адресную строку
        addr_actions = [
            {'type': 'class', 'class': 'android.widget.EditText'},
            {'type': 'text', 'text': 'Поиск', 'partial': True},
            {'type': 'text', 'text': 'Адрес', 'partial': True},
        ]
        if not self._click_with_fallback(addr_actions, timeout=8):
            return "Не найдена адресная строка"
        
        time.sleep(1)
        self.execute_automation([{'action': 'input_focused', 'text': query}])
        time.sleep(1)
        
        # Подтверждаем поиск
        confirm_actions = [
            {'type': 'text', 'text': 'Перейти', 'partial': True},
            {'type': 'text', 'text': 'Найти', 'partial': True},
            {'type': 'text', 'text': 'Поиск', 'partial': True},
            {'type': 'class', 'class': 'android.widget.ImageButton'},
        ]
        if not self._click_with_fallback(confirm_actions, timeout=4):
            self.execute_automation([{'action': 'press_enter'}])
        
        time.sleep(10)  # ждём загрузки результатов
        
        # Пробуем кликнуть по первой ссылке (разные варианты)
        link_actions = [
            {'type': 'class', 'class': 'android.widget.TextView'},
            {'type': 'class', 'class': 'android.view.View'},
            {'type': 'text', 'text': 'Открыть', 'partial': True},
            {'type': 'desc', 'desc': 'Перейти по ссылке'},
            {'type': 'class', 'class': 'android.widget.LinearLayout'},
        ]
        if not self._click_with_fallback(link_actions, timeout=8):
            # Если не нашли, пробуем кликнуть по центру экрана (иногда помогает)
            pass
        
        time.sleep(6)
        article_text = self.execute_automation([{'action': 'get_text'}])
        
        full_result = f"🔍 Поиск в браузере по запросу '{query}':\n\n{article_text[:1500]}"
        self.messages.append({"role": "system", "content": full_result})
        self.save_history()
        threading.Thread(target=self._call_api, daemon=True).start()
        return full_result

    def play_track_in_miui(self, track_name):
        if not self.android_bridge:
            return "Нет доступа к Android Bridge"
        
        self.android_bridge.openApp("Музыка")
        time.sleep(8)
        
        # ---- Шаг 1: Открыть поиск ----
        search_actions = [
            {'type': 'text', 'text': 'Поиск', 'partial': True},
            {'type': 'class', 'class': 'android.widget.ImageButton'},
            {'type': 'desc', 'desc': 'Поиск'},
            {'type': 'class', 'class': 'android.widget.Button'},
        ]
        if not self._click_with_fallback(search_actions, timeout=10):
            self._dump_screen_text("Музыка - не найден поиск")
            return "Не удалось открыть поиск"
        
        time.sleep(3)
        # ---- Шаг 2: Ввести название трека ----
        input_actions = [
            {'type': 'class', 'class': 'android.widget.EditText'},
            {'type': 'class', 'class': 'android.widget.TextView'},
        ]
        self._click_with_fallback(input_actions, timeout=5)
        time.sleep(1)
        self.execute_automation([{'action': 'input_focused', 'text': track_name}])
        time.sleep(1)
        
        # ---- Шаг 3: Нажать Enter или кнопку поиска ----
        confirm_actions = [
            {'type': 'text', 'text': 'Найти', 'partial': True},
            {'type': 'text', 'text': 'Поиск', 'partial': True},
            {'type': 'class', 'class': 'android.widget.ImageButton'},
        ]
        if not self._click_with_fallback(confirm_actions, timeout=3):
            self.execute_automation([{'action': 'press_enter'}])
        
        time.sleep(6)  # Ждём результаты
        
        # ---- Шаг 4: Кликнуть по первому треку ----
        track_actions = [
            {'type': 'class', 'class': 'android.widget.TextView'},
            {'type': 'class', 'class': 'android.widget.LinearLayout'},
            {'type': 'class', 'class': 'android.widget.RelativeLayout'},
        ]
        if not self._click_with_fallback(track_actions, timeout=8):
            self._dump_screen_text("Музыка - не найден трек")
            return "Не удалось выбрать трек"
        
        time.sleep(3)
        # ---- Шаг 5: Нажать Play ----
        play_actions = [
            {'type': 'text', 'text': 'Воспроизвести', 'partial': True},
            {'type': 'text', 'text': 'Play', 'partial': True},
            {'type': 'class', 'class': 'android.widget.ImageButton'},
            {'type': 'desc', 'desc': 'Воспроизвести'},
            {'type': 'class', 'class': 'android.widget.Button'},
        ]
        for attempt in range(3):
            if self._click_with_fallback(play_actions, timeout=5):
                break
            time.sleep(2)
        else:
            self._dump_screen_text("Музыка - не найдена кнопка Play")
            return "Не удалось воспроизвести трек"
        
        self.messages.append({"role": "system", "content": f"🎵 Трек '{track_name}' запущен"})
        self.save_history()
        threading.Thread(target=self._call_api, daemon=True).start()
        return "Трек запущен"

    def create_calendar_event(self, title, date=None):
        if not self.android_bridge:
            return "Нет доступа к Android Bridge"
        
        self.android_bridge.openApp("Календарь")
        time.sleep(8)
        
        # ---- Шаг 1: Кнопка создания ----
        create_actions = [
            {'type': 'text', 'text': 'Создать', 'partial': True},
            {'type': 'text', 'text': 'Добавить', 'partial': True},
            {'type': 'class', 'class': 'android.widget.ImageButton'},
            {'type': 'class', 'class': 'android.widget.FloatingActionButton'},
            {'type': 'class', 'class': 'android.widget.Button'},
            {'type': 'desc', 'desc': 'Создать событие'},
        ]
        if not self._click_with_fallback(create_actions, timeout=12):
            self._dump_screen_text("Календарь - не найдена кнопка создания")
            return "Не удалось создать событие"
        
        time.sleep(3)
        # ---- Шаг 2: Ввод заголовка ----
        input_actions = [
            {'type': 'class', 'class': 'android.widget.EditText'},
            {'type': 'text', 'text': 'Название', 'partial': True},
            {'type': 'text', 'text': 'Заголовок', 'partial': True},
        ]
        self._click_with_fallback(input_actions, timeout=5)
        time.sleep(1)
        self.execute_automation([{'action': 'input_focused', 'text': title}])
        time.sleep(2)
        
        # ---- Шаг 3: Сохранить ----
        save_actions = [
            {'type': 'text', 'text': 'Сохранить', 'partial': True},
            {'type': 'text', 'text': 'Готово', 'partial': True},
            {'type': 'text', 'text': 'ОК', 'partial': True},
            {'type': 'class', 'class': 'android.widget.Button'},
            {'type': 'desc', 'desc': 'Сохранить'},
        ]
        for attempt in range(3):
            if self._click_with_fallback(save_actions, timeout=8):
                break
            time.sleep(2)
        else:
            self._dump_screen_text("Календарь - не найдена кнопка сохранения")
            return "Не удалось сохранить событие"
        
        self.messages.append({"role": "system", "content": f"📅 Событие '{title}' создано"})
        self.save_history()
        threading.Thread(target=self._call_api, daemon=True).start()
        return "Событие создано"

    def create_text_file_in_downloads(self, filename, content):
        if self.android_bridge:
            downloads = self.android_bridge.getDownloadsPath()
            filepath = os.path.join(downloads, filename)
            try:
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(content)
                return f"✅ Файл '{filename}' создан в папке Загрузки."
            except Exception as e:
                return f"❌ Ошибка создания файла: {e}"
        else:
            return "Нет доступа к хранилищу."

    def _click_with_fallback(self, actions, timeout=10):
        start = time.time()
        while time.time() - start < timeout:
            for action in actions:
                if action['type'] == 'text':
                    result = self.execute_automation([{'action': 'click_by_text', 'text': action['text'], 'partial': action.get('partial', False)}])
                    if 'OK' in result:
                        time.sleep(0.5)
                        return True
                elif action['type'] == 'class':
                    result = self.execute_automation([{'action': 'click_by_class', 'class': action['class']}])
                    if 'OK' in result:
                        time.sleep(0.5)
                        return True
                elif action['type'] == 'desc':
                    result = self.execute_automation([{'action': 'click_by_content_desc', 'desc': action['desc']}])
                    if 'OK' in result:
                        time.sleep(0.5)
                        return True
                elif action['type'] == 'action' and action.get('action') == 'back':
                    self.execute_automation([{'action': 'back'}])
                    time.sleep(0.5)
                    return True
            time.sleep(0.5)
        return False

    def _wait_for_text_in_window(self, expected_text, timeout=8):
        start = time.time()
        while time.time() - start < timeout:
            window_text = self.execute_automation([{'action': 'get_text'}])
            if expected_text.lower() in window_text.lower():
                return True
            time.sleep(0.5)
        return False

    def _input_text_and_confirm(self, text, confirm_actions=None):
        self.execute_automation([{'action': 'click_by_class', 'class': 'android.widget.EditText'}])
        time.sleep(0.5)
        self.execute_automation([{'action': 'input_focused', 'text': text}])
        time.sleep(1)
        if confirm_actions:
            self._click_with_fallback(confirm_actions, timeout=3)
        else:
            self.execute_automation([{'action': 'press_enter'}])
        time.sleep(1)
    
    def _transcribe_audio(self, audio_bytes, filename):
        url = "https://bothub.chat/api/v2/openai/v1/audio/transcriptions"
        headers = {"Authorization": f"Bearer {self.api_key}"}
        
        # Определяем MIME-тип по расширению
        ext = filename.split('.')[-1].lower()
        mime_map = {
            'mp3': 'audio/mpeg',
            'm4a': 'audio/mp4',
            '3gp': 'audio/3gpp',
            'wav': 'audio/wav',
            'ogg': 'audio/ogg',
            'flac': 'audio/flac',
        }
        mime_type = mime_map.get(ext, 'audio/mpeg')
        
        if not audio_bytes:
            print("❌ _transcribe_audio: пустые байты")
            return None
        
        print(f"📤 Транскрипция: {filename}, размер {len(audio_bytes)} байт, MIME {mime_type}")
        
        files = {"file": (filename, audio_bytes, mime_type)}
        data = {"model": "whisper-1", "response_format": "text", "language": "ru"}
        
        try:
            response = requests.post(url, headers=headers, files=files, data=data, timeout=120)
            print(f"📡 Ответ сервера: статус {response.status_code}")
            
            if response.status_code == 200:
                content_type = response.headers.get('content-type', '')
                if 'application/json' in content_type:
                    result = response.json()
                    text = result.get("text", "").strip()
                else:
                    text = response.text.strip()
                
                if text:
                    print(f"✅ Транскрипция успешна: {text[:100]}...")
                    return text
                else:
                    print("⚠️ Транскрипция вернула пустую строку")
                    return None
            else:
                # Пытаемся извлечь сообщение об ошибке
                error_detail = response.text
                try:
                    error_json = response.json()
                    error_detail = error_json.get('error', {}).get('message', error_detail)
                except:
                    pass
                print(f"❌ Ошибка API ({response.status_code}): {error_detail}")
                return None
        except requests.exceptions.Timeout:
            print("❌ Таймаут при транскрипции (120 сек)")
            return None
        except Exception as e:
            print(f"❌ Исключение при транскрипции: {e}")
            return None

    def _play_media_file(self, filepath):
        if not self.android_bridge:
            self.messages.append({"role": "system", "content": "Нет доступа к Android Bridge для воспроизведения."})
            return False
        if isinstance(filepath, (list, tuple)):
            filepath = filepath[0]
        try:
            self.android_bridge.openFile(filepath)
            self.android_bridge.playOpen()
            return True
        except Exception as e:
            self.log(f"Ошибка воспроизведения медиа: {e}", "red", "system")
            return False

    def generate_image(self, prompt):
        model = "gemini-3-pro-image-preview"
        payload = {
            "model": model,
            "messages": [{"role": "user", "content": prompt}],
            "image_config": {"aspect_ratio": "1:1", "image_size": "2K"},
            "response_modalities": ["IMAGE", "TEXT"]
        }

        try:
            response = requests.post(API_URL, headers=self.headers, json=payload, timeout=120)
            if response.status_code != 200:
                error_msg = f"❌ Ошибка API: {response.status_code}"
                self.messages.append({"role": "system", "content": error_msg})
                self.save_history()
                self._call_api()
                return False

            data = response.json()
            message = data["choices"][0]["message"]

            # Проверяем наличие изображений в новом формате
            if "images" in message and message["images"]:
                image_url = message["images"][0]["image_url"]["url"]
                if image_url.startswith("data:"):
                    base64_data = image_url.split(",", 1)[1]
                    if self.android_bridge:
                        self.android_bridge.saveBase64Image(base64_data, "png")
                        self.messages.append({"role": "system", "content": "🎨 Изображение сгенерировано и сохранено в галерею."})
                    else:
                        self.messages.append({"role": "system", "content": "🎨 Изображение получено, но не сохранено (нет bridge)."})
                    self.save_history()
                    self._call_api()
                    return True

            # Проверяем старый формат inline_data
            if "inline_data" in message and message["inline_data"].get("data"):
                base64_data = message["inline_data"]["data"]
                if self.android_bridge:
                    self.android_bridge.saveBase64Image(base64_data, "png")
                    self.messages.append({"role": "system", "content": "🎨 Изображение сгенерировано и сохранено в галерею."})
                else:
                    self.messages.append({"role": "system", "content": "🎨 Изображение получено, но не сохранено (нет bridge)."})
                self.save_history()
                self._call_api()
                return True

            self.messages.append({"role": "system", "content": f"⚠️ Не удалось извлечь изображение. Ответ:\n{message.get('content', 'Пусто')}"})
            self.save_history()
            self._call_api()
            return True

        except Exception as e:
            error_msg = f"❌ Ошибка при генерации: {e}"
            self.messages.append({"role": "system", "content": error_msg})
            self.save_history()
            self._call_api()
            return False
    
    def search_wikipedia(self, query):
        """Поиск в Википедии с обработкой ошибок"""
        try:
            print(f"[WIKI] Ищу: {query}")
            wikipedia.set_lang("ru")
            results = wikipedia.search(query, results=2)
            if not results:
                print("[WIKI] Нет результатов поиска")
                return None
            output = []
            for title in results:
                try:
                    page = wikipedia.page(title)
                    summary = page.summary[:300] + "..." if len(page.summary) > 300 else page.summary
                    output.append(f"**{page.title}**\n{summary}\n[Читать]({page.url})")
                    print(f"[WIKI] Найдено: {page.title}")
                except Exception as e:
                    print(f"[WIKI] Ошибка при получении страницы {title}: {e}")
                    continue
            return output if output else None
        except Exception as e:
            print(f"[WIKI] Ошибка поиска: {e}")
            return None

    def search_news(self, query):
        news_sources = {
            "ria": {  # РИА Новости
                "url": "https://ria.ru/search/",
                "type": "ria",
                "params": {"query": query}
            },
            "cnews": {  # Новости высоких технологий
                "url": "https://www.cnews.ru/search/",
                "type": "cnews",
                "params": {"text": query}
            },
            "ferra": {  # Наука и технологии
                "url": "https://www.ferra.ru/search/",
                "type": "ferra",
                "params": {"q": query}
            },
            "lenta": {  # Lenta.ru: Наука и техника
                "url": "https://lenta.ru/rss/news",
                "type": "rss",
                "feed_url": "https://lenta.ru/rss/news"
            },
            "kommersant": {  # Коммерсантъ
                "url": "https://www.kommersant.ru/RSS/news.xml",
                "type": "rss",
                "feed_url": "https://www.kommersant.ru/RSS/news.xml"
            }
        }

        def parse_rss(feed_url):
            try:
                import feedparser
                response = requests.get(feed_url, timeout=10)
                if response.status_code == 200:
                    feed = feedparser.parse(response.content)
                    output = []
                    for entry in feed.entries[:5]:
                        title = entry.get('title', 'Без названия')
                        link = entry.get('link', '#')
                        output.append(f"**{title}**\n[Читать]({link})")
                    return output if output else None
                else:
                    return None
            except Exception as e:
                print(f"Ошибка парсинга RSS {feed_url}: {e}")
                return None

        for source_name, source_data in news_sources.items():
            try:
                if source_data["type"] == "ria":
                    url = source_data["url"]
                    params = source_data["params"]
                    response = requests.get(url, params=params, timeout=10)
                    if response.status_code == 200:
                        soup = BeautifulSoup(response.text, 'html.parser')
                        output = []
                        for link in soup.find_all('a', class_='list-item__title'):
                            title = link.get_text(strip=True)
                            href = link.get('href')
                            if href and not href.startswith('http'):
                                href = 'https://ria.ru' + href
                            output.append(f"**{title}**\n[Читать]({href})")
                            if len(output) >= 3:
                                break
                        if output:
                            return output
                elif source_data["type"] == "cnews":
                    url = source_data["url"]
                    params = source_data["params"]
                    response = requests.get(url, params=params, timeout=10)
                    if response.status_code == 200:
                        soup = BeautifulSoup(response.text, 'html.parser')
                        output = []
                        for link in soup.find_all('a', class_='news-item__title'):
                            title = link.get_text(strip=True)
                            href = link.get('href')
                            if href:
                                output.append(f"**{title}**\n[Читать]({href})")
                                if len(output) >= 3:
                                    break
                        if output:
                            return output
                elif source_data["type"] == "ferra":
                    url = source_data["url"]
                    params = source_data["params"]
                    response = requests.get(url, params=params, timeout=10)
                    if response.status_code == 200:
                        soup = BeautifulSoup(response.text, 'html.parser')
                        output = []
                        for link in soup.find_all('a', class_='entry-title'):
                            title = link.get_text(strip=True)
                            href = link.get('href')
                            if href:
                                output.append(f"**{title}**\n[Читать]({href})")
                                if len(output) >= 3:
                                    break
                        if output:
                            return output
                elif source_data["type"] == "rss":
                    rss_output = parse_rss(source_data["feed_url"])
                    if rss_output:
                        filtered = []
                        for item in rss_output:
                            if query.lower() in item.lower():
                                filtered.append(item)
                        if filtered:
                            return filtered[:3]
                        else:
                            return rss_output[:3]  # Если по запросу не отфильтровалось, показываем последние новости
            except Exception as e:
                print(f"Ошибка поиска в {source_name}: {e}")
                continue

        return None

    def process_search(self, query, search_type="general"):
        """Универсальный поиск с гарантированным ответом Гали"""
        print(f"[SEARCH] Запрос: {query}, тип: {search_type}")

        if not query or not query.strip():
            self.messages.append({"role": "system", "content": "Пустой поисковый запрос."})
            self.save_history()
            threading.Thread(target=self._call_api, daemon=True).start()
        
        # --- НОВОСТИ ---
        if search_type == "news":
            try:
                results = self.search_news(query)
            except Exception as e:
                print(f"[NEWS] Ошибка: {e}")
                results = None

            if results:
                output = f"**Новости по запросу \"{query}\":**\n\n" + "\n\n---\n\n".join(results)
                self.messages.append({
                    "role": "system",
                    "content": f"Пользователь искал новости '{query}'. Вот результаты:\n{output}"
                })
                self.save_history()
            else:
                url = f"https://www.google.com/search?q={urllib.parse.quote(query)}+новости&tbm=nws"
                if self.android_bridge:
                    self.android_bridge.openUrl(url)
                    self.messages.append({
                        "role": "system",
                        "content": f"Новостей не найдено, открываю поиск Google по запросу '{query}'."
                    })
                else:
                    self.messages.append({
                        "role": "system",
                        "content": f"Не удалось открыть поиск: нет доступа к браузеру."
                    })
                self.save_history()
            threading.Thread(target=self._call_api, daemon=True).start()
            return

        # --- ВИКИПЕДИЯ ---
        if search_type == "wiki":
            try:
                results = self.search_wikipedia(query)
            except Exception as e:
                print(f"[WIKI] Ошибка при вызове search_wikipedia: {e}")
                results = None
            
            if results:
                output = f"**Википедия: {query}**\n\n" + "\n\n---\n\n".join(results)
                self.messages.append({
                    "role": "system",
                    "content": f"Пользователь искал в Википедии '{query}'. Вот результаты:\n{output}"
                })
                print("[WIKI] Найдены результаты, добавлено сообщение в историю")
            else:
                # Открываем браузер с поиском в Википедии
                url = f"https://ru.wikipedia.org/wiki/{urllib.parse.quote(query)}"
                if self.android_bridge:
                    self.android_bridge.openUrl(url)
                    self.messages.append({
                        "role": "system",
                        "content": f"Открываю статью в Википедии по запросу '{query}'."
                    })
                    print("[WIKI] Открываю браузер со статьёй Википедии")
                else:
                    self.messages.append({
                        "role": "system",
                        "content": f"Не удалось открыть Википедию: нет доступа к браузеру."
                    })
                    print("[WIKI] Ошибка: нет android_bridge")
            self.save_history()
            threading.Thread(target=self._call_api, daemon=True).start()
            return

        # --- ОБЩИЙ ПОИСК ---
        if search_type == "general":
            results = self.search_searxng(query)
            if results:
                self.last_search_results = []
                for res in results:
                    url_match = re.search(r'\[Ссылка\]\(([^)]+)\)', res)
                    if url_match:
                        self.last_search_results.append(url_match.group(1))
                output = f"**Результаты поиска по запросу \"{query}\":**\n\n" + "\n\n---\n\n".join(results)
                self.messages.append({
                    "role": "system",
                    "content": f"Пользователь искал '{query}'. Вот результаты:\n{output}"
                })
            else:
                url = f"https://www.google.com/search?q={urllib.parse.quote(query)}"
                if self.android_bridge:
                    self.android_bridge.openUrl(url)
                    self.messages.append({
                        "role": "system",
                        "content": f"Не удалось получить результаты поиска, открываю Google по запросу '{query}' в браузере."
                    })
                else:
                    self.messages.append({
                        "role": "system",
                        "content": f"Не удалось выполнить поиск: нет доступа к браузеру."
                    })
            self.save_history()
            threading.Thread(target=self._call_api, daemon=True).start()
            return

    def process_image(self, base64_image, description=""):
            def _process():
                try:
                    # Проверяем размер base64 (приблизительно)
                    size_mb = len(base64_image) / (1024 * 1024)
                    if size_mb > 10:  # если больше 10 МБ, предупреждаем
                        print(f"⚠️ Изображение большое: {size_mb:.1f} МБ, обработка может занять время")
                    
                    print(f"📷 Получено изображение: {len(base64_image)} байт")
                    image_content = {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/jpeg;base64,{base64_image}"
                        }
                    }
                    text_content = {
                        "type": "text",
                        "text": description if description else "Опиши что на этом изображении. Будь краткой и полезной."
                    }
                    self.messages.append({
                        "role": "user",
                        "content": [text_content, image_content]
                    })
                    self.save_history()
                    # Вызываем API в этом же потоке (он уже фоновый)
                    self._call_api()
                except Exception as e:
                    print(f"❌ Ошибка обработки изображения: {e}")
                    if self.android_bridge:
                        self.android_bridge.addAssistantMessage(f"❌ Ошибка при обработке изображения: {e}")
                    # Добавляем системное сообщение, чтобы Галя могла ответить
                    self.messages.append({
                        "role": "system",
                        "content": f"Ошибка при обработке изображения: {e}"
                    })
                    self.save_history()
                    # Вызываем API, чтобы Галя отреагировала
                    threading.Thread(target=self._call_api, daemon=True).start()

            # Запускаем обработку в фоне
            threading.Thread(target=_process, daemon=True).start()

    def find_file_in_storage(self, filename):
        if not self.android_bridge:
            self.messages.append({"role": "system", "content": "Не могу искать файлы: нет доступа к хранилищу."})
            self.save_history()
            threading.Thread(target=self._call_api, daemon=True).start()
            return

        try:
            results = self.android_bridge.searchFiles(filename)
            if results:
                output = f"Найдены файлы по запросу '{filename}':\n" + "\n".join(results)
                self.messages.append({"role": "system", "content": output})
                if self.android_bridge:
                    self.android_bridge.playDone()
            else:
                self.messages.append({"role": "system", "content": f"Файлы '{filename}' не найдены."})
                if self.android_bridge:
                    self.android_bridge.playError()
            self.save_history()
            threading.Thread(target=self._call_api, daemon=True).start()
        except Exception as e:
            self.messages.append({"role": "system", "content": f"Ошибка при поиске файлов: {e}"})
            if self.android_bridge:
                self.android_bridge.playError()
            self.save_history()
            threading.Thread(target=self._call_api, daemon=True).start()            

    def open_file_by_name(self, filename):
        if not self.android_bridge:
            self.messages.append({"role": "system", "content": "Не удалось открыть файл: нет доступа к хранилищу."})
            self.save_history()
            threading.Thread(target=self._call_api, daemon=True).start()
            return

        try:
            results = self.android_bridge.searchFiles(filename)
        except Exception as e:
            self.messages.append({"role": "system", "content": f"Ошибка поиска файла '{filename}': {e}"})
            if self.android_bridge:
                self.android_bridge.playError()
            self.save_history()
            threading.Thread(target=self._call_api, daemon=True).start()
            return

        if not results:
            self.messages.append({"role": "system", "content": f"Файл '{filename}' не найден."})
            if self.android_bridge:
                self.android_bridge.playError()
            self.save_history()
            threading.Thread(target=self._call_api, daemon=True).start()
            return

        if len(results) == 1:
            try:
                self.android_bridge.openFile(results[0])
                self.messages.append({"role": "system", "content": f"Открываю файл: {results[0]}"})
                if self.android_bridge:
                    self.android_bridge.playOpen()
                    self.android_bridge.playDone()
            except Exception as e:
                self.messages.append({"role": "system", "content": f"Не удалось открыть файл {results[0]}: {e}"})
                if self.android_bridge:
                    self.android_bridge.playError()
            self.save_history()
            threading.Thread(target=self._call_api, daemon=True).start()
        else:
            output = "Найдено несколько файлов:\n" + "\n".join(results)
            self.messages.append({"role": "system", "content": output})
            if self.android_bridge:
                self.android_bridge.playError()  # или playDone, но пусть будет ошибка
            self.save_history()
            threading.Thread(target=self._call_api, daemon=True).start()
    
    def process_uploaded_text(self, filename, content):
        ext = os.path.splitext(filename)[1].lower()
        text_extensions = ['.txt', '.py', '.json', '.xml', '.html', '.css', '.js', '.md', '.csv']

        # Аудиофайлы — транскрибируем
        if ext in ['.mp3', '.wav', '.ogg', '.flac', '.m4a', '.3gp']:
            print(f"[AUDIO] Обнаружен аудиофайл: {filename}, расширение {ext}, размер данных: {len(content) if isinstance(content, bytes) else 'путь'}")
            self.messages.append({
                "role": "system",
                "content": f"🎵 Получен аудиофайл: {filename}. Распознаю речь..."
            })
            self.save_history()
            
            def process_audio():
                # Убедимся, что content — это байты
                if isinstance(content, str):
                    # Если передан путь, читаем файл
                    try:
                        with open(content, 'rb') as f:
                            audio_bytes = f.read()
                    except Exception as e:
                        self.messages.append({
                            "role": "system",
                            "content": f"❌ Не удалось прочитать аудиофайл с диска: {e}"
                        })
                        self.save_history()
                        self._call_api()
                        return
                else:
                    audio_bytes = content
                
                if not audio_bytes:
                    self.messages.append({
                        "role": "system",
                        "content": f"❌ Аудиофайл {filename} пуст или повреждён."
                    })
                    self.save_history()
                    self._call_api()
                    return
                
                text = self._transcribe_audio(audio_bytes, filename)
                if text:
                    self.messages.append({
                        "role": "system",
                        "content": f"📝 Транскрипция аудиофайла {filename}:\n```\n{text}\n```"
                    })
                    if self.android_bridge:
                        self.android_bridge.playDone()
                else:
                    self.messages.append({
                        "role": "system",
                        "content": f"❌ Не удалось распознать речь в файле {filename}. Проверьте формат (должен быть MP3, M4A, WAV, OGG, FLAC, 3GP) и качество записи."
                    })
                    if self.android_bridge:
                        self.android_bridge.playError()
                self.save_history()
                threading.Thread(target=self._call_api, daemon=True).start()
            
            threading.Thread(target=process_audio, daemon=True).start()
            return
        
        # Видеофайлы — не транскрибируем, а только открываем как медиа
        if ext in ['.mp4', '.avi', '.mov', '.mkv', '.webm']:
            self.messages.append({
                "role": "system",
                "content": f"🎬 Получен видеофайл: {filename}. Я не могу распознать речь из видео, но могу открыть его плеером."
            })
            self.save_history()
            threading.Thread(target=self._call_api, daemon=True).start()
            return
        
        # Медиафайлы (изображения) — оставляем как есть, только уведомление
        if ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
            self.messages.append({
                "role": "system",
                "content": f"📷 Получен медиафайл: {filename} (тип {ext}). Я не могу прочитать его содержимое, но могу выполнить другие действия."
            })
            self.save_history()
            threading.Thread(target=self._call_api, daemon=True).start()
            return

        # PDF
        if ext == '.pdf':
            if PdfReader is None:
                self.messages.append({
                    "role": "system",
                    "content": f"❌ Для чтения PDF требуется библиотека pypdf. Файл {filename} не может быть обработан."
                })
                self.save_history()
                threading.Thread(target=self._call_api, daemon=True).start()
                return
            try:
                if isinstance(content, bytes):
                    import io
                    reader = PdfReader(io.BytesIO(content))
                else:
                    with open(content, 'rb') as f:
                        reader = PdfReader(f)
                text = ""
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                if not text.strip():
                    text = "[Текст не найден в PDF]"
                self.messages.append({
                    "role": "system",
                    "content": f"📄 Пользователь загрузил PDF файл {filename}. Содержимое:\n```\n{text}\n```"
                })
            except Exception as e:
                self.messages.append({
                    "role": "system",
                    "content": f"❌ Не удалось прочитать PDF {filename}: {e}"
                })
            self.save_history()
            threading.Thread(target=self._call_api, daemon=True).start()
            return

        # DOCX
        if ext == '.docx':
            self._process_docx(filename, content)
            return

        # Старый DOC
        if ext == '.doc':
            self._process_doc(filename, content)
            return

        # Текстовые файлы (читаем полностью, но сначала проверяем бинарность)
        if ext in text_extensions or ext == '':
            if isinstance(content, bytes):
                if is_likely_binary(content):
                    self.messages.append({
                        "role": "system",
                        "content": f"📁 Получен файл {filename}, но он выглядит как бинарный. Я не могу прочитать его содержимое."
                    })
                    self.save_history()
                    threading.Thread(target=self._call_api, daemon=True).start()
                    return
                try:
                    content = content.decode('utf-8')
                except UnicodeDecodeError:
                    try:
                        content = content.decode('cp1251')
                    except UnicodeDecodeError:
                        content = content.decode('latin-1', errors='replace')
            self.messages.append({
                "role": "system",
                "content": f"📄 Пользователь загрузил файл {filename}. Содержимое:\n```\n{content}\n```"
            })
            self.save_history()
            threading.Thread(target=self._call_api, daemon=True).start()
            return

        # Неизвестный тип
        self.messages.append({
            "role": "system",
            "content": f"📁 Получен файл: {filename} (тип {ext}). Тип не распознан. Я не могу обработать его содержимое."
        })
        self.save_history()
        threading.Thread(target=self._call_api, daemon=True).start()

    def _process_doc(self, filename, content):
        encodings = ['utf-8', 'cp1251', 'koi8-r', 'iso-8859-5', 'cp866', 'mac-cyrillic', 'utf-16', 'utf-16-le', 'utf-16-be']

        def decode_text(data_bytes):
            for enc in encodings:
                try:
                    return data_bytes.decode(enc, errors='ignore').strip()
                except:
                    continue
            return data_bytes.decode('utf-8', errors='ignore')

        try:
            if isinstance(content, bytes):
                raw = content
            else:
                with open(content, 'rb') as f:
                    raw = f.read()
            text = decode_text(raw)
            if text:
                self._finalize_processing(filename, text)
                return
        except Exception as e:
            self.log(f"Ошибка при чтении DOC как текст: {e}", "yellow", "system")

        try:
            from olefile import OleFileIO
            import io

            if isinstance(content, bytes):
                ole = OleFileIO(io.BytesIO(content))
            else:
                ole = OleFileIO(content)

            all_text = []

            def extract_from_storage(storage):
                for entry in storage.listdir():
                    # Поток
                    try:
                        stream = storage.openstream(entry)
                        data = stream.read()
                        if data:
                            text = decode_text(data)
                            if text:
                                all_text.append(text)
                    except:
                        pass
                    try:
                        if storage.exists(entry) and storage.get_type(entry) == 1:
                            sub = storage.openstorage(entry)
                            extract_from_storage(sub)
                    except:
                        pass

            extract_from_storage(ole)
            ole.close()

            if all_text:
                full_text = "\n".join(all_text)
                self._finalize_processing(filename, full_text)
                return
            else:
                raise Exception("olefile не извлёк текст")
        except Exception as e:
            self.log(f"Ошибка olefile: {e}", "red", "system")
            self.messages.append({"role": "system", "content": f"❌ Не удалось прочитать документ {filename}. Формат может быть повреждён или защищён."})
            self.save_history()
            threading.Thread(target=self._call_api, daemon=True).start()

    def _finalize_processing(self, filename, text_content):
        if not text_content.strip():
            text_content = "[Текст не найден в документе]"
        
        self.messages.append({
            "role": "system",
            "content": f"📄 Пользователь загрузил документ {filename}. Содержимое:\n```\n{text_content[:2000]}\n```"
        })
        self.save_history()
        self.log(f"DOC файл {filename} обработан, текст добавлен в историю ({len(text_content)} символов)", "green")
        threading.Thread(target=self._call_api, daemon=True).start()
    
    def _process_docx(self, filename, content):
        if Document is None:
            self.messages.append({
                "role": "system",
                "content": f"❌ Для чтения DOCX требуется библиотека python-docx. Файл {filename} не может быть обработан."
            })
            self.save_history()
            threading.Thread(target=self._call_api, daemon=True).start()
            return

        try:
            if isinstance(content, bytes):
                import io
                doc = Document(io.BytesIO(content))
            else:
                doc = Document(content)

            text = "\n".join([para.text for para in doc.paragraphs])
            if not text.strip():
                text = "[Текст не найден в документе]"

            self.messages.append({
                "role": "system",
                "content": f"📄 Пользователь загрузил документ {filename}. Содержимое:\n```\n{text}\n```"
            })
            self.save_history()
            self.log(f"DOCX файл {filename} обработан, текст добавлен в историю ({len(text)} символов)", "green")
        except Exception as e:
            self.messages.append({
                "role": "system",
                "content": f"❌ Не удалось прочитать документ {filename}: {e}"
            })
            self.save_history()
            self.log(f"Ошибка DOCX: {e}", "red")
        threading.Thread(target=self._call_api, daemon=True).start()

    def _process_pdf(self, filename, content):
        try:
            from pypdf import PdfReader
        except ImportError:
            self.messages.append({"role": "system", "content": f"❌ Для чтения PDF требуется библиотека pypdf. Файл {filename} не может быть обработан."})
            self.save_history()
            threading.Thread(target=self._call_api, daemon=True).start()
            return

        try:
            if isinstance(content, bytes):
                import io
                reader = PdfReader(io.BytesIO(content))
            else:
                if not os.path.exists(content):
                    raise FileNotFoundError(f"Файл {content} не найден")
                with open(content, 'rb') as f:
                    reader = PdfReader(f)

            full_text = []
            for page_num, page in enumerate(reader.pages, 1):
                page_text = ""
                # Стандартное извлечение
                try:
                    page_text = page.extract_text()
                except:
                    pass
                # Альтернативный режим (если поддерживается)
                if not page_text or len(page_text.strip()) < 10:
                    try:
                        page_text = page.extract_text(extraction_mode="layout")
                    except:
                        pass
                # Аннотации
                if not page_text or len(page_text.strip()) < 10:
                    try:
                        if hasattr(page, 'annotations') and page.annotations:
                            for annot in page.annotations:
                                if annot.get("/Contents"):
                                    page_text += annot.get("/Contents", "") + " "
                    except:
                        pass
                if page_text and page_text.strip():
                    full_text.append(page_text.strip())
                else:
                    full_text.append(f"[Страница {page_num} не содержит текста]")

            result_text = "\n\n".join(full_text)
            if not result_text.strip():
                result_text = "[PDF не содержит текста. Возможно, это скан-копия.]"

            self._finalize_processing(filename, result_text)

        except FileNotFoundError:
            self.messages.append({"role": "system", "content": f"❌ Файл {filename} не найден."})
            self.save_history()
            self.log(f"Ошибка PDF: файл не найден {filename}", "red")
            threading.Thread(target=self._call_api, daemon=True).start()
        except Exception as e:
            self.messages.append({"role": "system", "content": f"❌ Не удалось прочитать PDF {filename}: {str(e)}"})
            self.save_history()
            self.log(f"Ошибка PDF: {e}", "red")
            threading.Thread(target=self._call_api, daemon=True).start()

    def read_url(self, url):
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        try:
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            response.encoding = response.apparent_encoding or 'utf-8'
            soup = BeautifulSoup(response.text, 'html.parser')
            for script in soup(["script", "style"]):
                script.decompose()
            text = soup.get_text(separator='\n', strip=True)
            if len(text) > 10000:
                text = text[:10000] + "...\n[Текст обрезан]"
            if text:
                self.messages.append({
                    "role": "system",
                    "content": f"Содержимое {url}:\n{text}"
                })
                self.save_history()
                threading.Thread(target=self._call_api, daemon=True).start()

        except Exception as e:
            pass

    def open_bookmark(self, name):
        bookmarks = load_bookmarks()
        name_lower = name.lower()
        for key, url in bookmarks.items():
            if name_lower in key.lower() or key.lower() in name_lower:
                if self.android_bridge:
                    self.android_bridge.openUrl(url)
                    self.android_bridge.playOpen()
                    self.messages.append({"role": "system", "content": f"Открываю закладку {key}: {url}"})
                    self.save_history()
                    threading.Thread(target=self._call_api, daemon=True).start()
                return
        self.messages.append({"role": "system", "content": f"Закладка '{name}' не найдена."})
        self.save_history()
        threading.Thread(target=self._call_api, daemon=True).start()

    def add_task(self, task_text):
        tasks = self.load_tasks()
        tasks.append({"id": len(tasks) + 1, "task": task_text, "done": False, "created": datetime.now().isoformat()})
        self.save_tasks(tasks)
        self.messages.append({"role": "system", "content": f"✅ Задача добавлена: {task_text}"})
        self.save_history()
        threading.Thread(target=self._call_api, daemon=True).start()

    def list_tasks(self):
        tasks = self.load_tasks()
        if not tasks:
            output = "Список задач пуст"
        else:
            output = "Твои задачи:\n"
            for t in tasks:
                status = "✅" if t["done"] else "⏳"
                output += f"{status} [{t['id']}] {t['task']}\n"
        self.messages.append({"role": "system", "content": output})
        self.save_history()
        threading.Thread(target=self._call_api, daemon=True).start()
        return output

    def complete_task(self, task_id):
        tasks = self.load_tasks()
        for t in tasks:
            if t["id"] == task_id:
                t["done"] = True
                self.save_tasks(tasks)
                self.messages.append({"role": "system", "content": f"✅ Задача {task_id} отмечена выполненной."})
                self.save_history()
                threading.Thread(target=self._call_api, daemon=True).start()
                return

    def read_file_from_disk(self, filepath):
        expanded = os.path.expandvars(filepath)
        if not os.path.exists(expanded):
            if self.android_bridge:
                base_name = os.path.basename(expanded)
                found_paths = self.android_bridge.searchFiles(base_name)
                if found_paths:
                    expanded = found_paths[0]
                else:
                    self.log(f"Файл не найден: {expanded}", "red", "system")
                    self.messages.append({"role": "system", "content": f"Файл {expanded} не найден."})
                    self.save_history()
                    threading.Thread(target=self._call_api, daemon=True).start()
                    return

        ext = os.path.splitext(expanded)[1].lower()

        # DOCX
        if ext == '.docx':
            self._process_docx(filename=os.path.basename(expanded), content=expanded)
            return

        # Старый DOC
        if ext == '.doc':
            self._process_doc(filename=os.path.basename(expanded), content=expanded)
            return

        # PDF
        if ext == '.pdf':
            self._process_pdf(filename=os.path.basename(expanded), content=expanded)
            return

        # Для остальных файлов пробуем читать как текст
        file_size = os.path.getsize(expanded)
        if file_size > 10 * 1024 * 1024:
            self.log(f"Файл слишком большой ({file_size} байт).", "red", "system")
            self.messages.append({"role": "system", "content": f"Файл {expanded} превышает 10 МБ и не может быть прочитан целиком."})
            self.save_history()
            threading.Thread(target=self._call_api, daemon=True).start()
            return

        # Проверка бинарности
        with open(expanded, 'rb') as f:
            raw_bytes = f.read(1024)
            f.seek(0)
            if is_likely_binary(raw_bytes):
                self.log(f"Файл {expanded} определён как бинарный, не читаем.", "yellow", "system")
                self.messages.append({"role": "system", "content": f"Файл {expanded} является бинарным и не может быть прочитан как текст."})
                self.save_history()
                threading.Thread(target=self._call_api, daemon=True).start()
                return

            # Читаем как текст
            content = None
            for enc in ['utf-8', 'cp1251', 'utf-8-sig', 'latin-1']:
                try:
                    with open(expanded, 'r', encoding=enc) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue

        if content is None:
            self.log(f"Файл {expanded} не является текстовым.", "red", "system")
            self.messages.append({"role": "system", "content": f"Файл {expanded} не является текстовым."})
            self.save_history()
            threading.Thread(target=self._call_api, daemon=True).start()
            return

        self.log(f"Содержимое файла {expanded} прочитано.", "white", "system")
        self.messages.append({
            "role": "system",
            "content": f"Файл {expanded} прочитан (размер {file_size} байт). Содержимое:\n```\n{content}\n```"
        })
        self.save_history()
        threading.Thread(target=self._call_api, daemon=True).start()

    def write_file_to_disk(self, filepath, content):
        # Если путь абсолютный, используем его
        if os.path.isabs(filepath) or filepath.startswith('/'):
            expanded = filepath
        else:
            # Если путь относительный, сохраняем в папку Загрузки
            if self.android_bridge:
                downloads = self.android_bridge.getDownloadsPath()
                expanded = os.path.join(downloads, filepath)
            else:
                expanded = filepath
        expanded = os.path.expandvars(expanded)
        try:
            os.makedirs(os.path.dirname(expanded), exist_ok=True)
            with open(expanded, 'w', encoding='utf-8') as f:
                f.write(content)
            self.log(f"Файл сохранён: {expanded}", "green", "system")
            self.messages.append({"role": "system", "content": f"Файл {expanded} успешно записан."})
        except Exception as e:
            self.log(f"Ошибка записи файла: {e}", "red", "system")
            self.messages.append({"role": "system", "content": f"Не удалось записать файл {expanded}: {e}"})
            self.save_history()
            threading.Thread(target=self._call_api, daemon=True).start()

    def _determine_effort(self, user_input):
        low_effort_keywords = ['привет', 'пока', 'как дела', 'что делаешь', 'спасибо', 'ок']
        high_effort_keywords = ['код', 'программа', 'напиши', 'объясни', 'почему', 'как работает', 'найди ошибку']
        text = user_input.lower()
        if any(kw in text for kw in high_effort_keywords):
            return "high"
        if any(kw in text for kw in low_effort_keywords):
            return "low"
        return "medium"

    def _call_api(self):
        if self._api_call_in_progress:
            print("⚠️ API вызов уже в процессе, пропускаем (защита от цикла)")
            return
        
        self._api_call_in_progress = True
        now = time.time()
        if hasattr(self, '_last_api_call') and now - self._last_api_call < 2:
            print("⚠️ Слишком частые вызовы API, пропускаем")
            return
        self._last_api_call = now

        self._trim_history_by_size()
        
        now = time.time()
        if not hasattr(self, '_call_times'):
                self._call_times = []
        self._call_times = [t for t in self._call_times if now - t < 60]  # очистить старые
        if len(self._call_times) >= 5:
                print("⚠️ Лимит вызовов API (5 в минуту), пропускаем")
                return
        self._call_times.append(now)

        if len(self.messages) > 21:
            self.messages = [self.messages[0]] + self.messages[-20:]
        last_user_msg = ""
        for msg in reversed(self.messages):
            if msg["role"] == "user":
                if isinstance(msg["content"], str):
                    last_user_msg = msg["content"]
                elif isinstance(msg["content"], list):
                    for item in msg["content"]:
                        if item.get("type") == "text":
                            last_user_msg = item["text"]
                            break
                break
        effort = self._determine_effort(last_user_msg)
        headers = self.headers.copy()
        messages_with_cache = []
        for i, msg in enumerate(self.messages):
            if i == 0 and msg["role"] == "system":
                if isinstance(msg["content"], str):
                    messages_with_cache.append({
                        "role": "system",
                        "content": [{"type": "text", "text": msg["content"], "cache_control": {"type": "ephemeral"}}]
                    })
                else:
                    messages_with_cache.append(msg)
            elif msg["role"] == "user":
                if isinstance(msg["content"], str):
                    messages_with_cache.append({
                        "role": "user",
                        "content": [{"type": "text", "text": msg["content"]}]
                    })
                else:
                    messages_with_cache.append(msg)
            else:
                messages_with_cache.append(msg)
        payload = {
            "model": self.model,
            "messages": messages_with_cache,
            "temperature": 0.7,
            "max_tokens": 5000,
            "thinking": {"type": "adaptive", "effort": effort}
        }
        try:
            response = session.post(API_URL, headers=headers, json=payload, timeout=120)
            if response.status_code == 200:
                data = response.json()
                if data.get("choices") and len(data["choices"]) > 0:
                    reply = data["choices"][0]["message"]["content"]
                    if isinstance(reply, list):
                        reply = "\n".join([item.get("text", "") for item in reply if item.get("type") == "text"])
                else:
                    reply = "Пустой ответ от API"
            elif response.status_code == 400:
                error_data = response.json() if response.text else {}
                error_msg = error_data.get('error', {}).get('message', '')
                if 'too long' in error_msg.lower() or 'maximum' in error_msg.lower():
                    self.log("Слишком длинный промпт, очищаю историю", "red")
                    system_msg = self.messages[0] if self.messages and self.messages[0]["role"] == "system" else None
                    self.messages = [system_msg] if system_msg else []
                    self.save_history()
                    reply = "История была слишком длинной. Я очистила её. Пожалуйста, повторите запрос."
                    self.messages.append({"role": "assistant", "content": reply})
                    self.save_history()
                    if self.android_bridge:
                        self.android_bridge.addAssistantMessage(reply)
                    self._api_call_in_progress = False
                    return
            else:
                reply = f"Ошибка API: {response.status_code}"
                print(f"❌ Ошибка API: {response.status_code} - {response.text}")
        except Exception as e:
            reply = f"Не удалось соединиться: {e}"
            print(f"❌ Исключение: {e}")
        if reply is None or reply.strip() == "":
            reply = "Не удалось получить ответ."
        self.messages.append({"role": "assistant", "content": reply})
        self.save_history()
        self._process_commands(reply)

        if self.android_bridge:
            self.android_bridge.addAssistantMessage(reply)

        self._api_call_in_progress = False
        
        return reply

    def _process_commands(self, text):
        if not isinstance(text, str):
            text = ""
        threading.Thread(target=self._execute_commands, args=(text,), daemon=True).start()

    def _execute_commands(self, text):
        executed = set()
        need_api_call = False

        def mark_executed(cmd_type, param=""):
            key = f"{cmd_type}:{param}"
            if key in executed:
                return False
            executed.add(key)
            return True

        # --- ЧТЕНИЕ URL ---
        for url in re.findall(r'\[READ_URL\](.*?)\[/READ_URL\]', text, re.DOTALL):
            url = url.strip()
            if mark_executed("READ_URL", url):
                self.read_url(url)
                need_api_call = True

        # --- ЗАКЛАДКИ ---
        for name in re.findall(r'\[BOOKMARK\](.*?)\[/BOOKMARK\]', text, re.DOTALL):
            name = name.strip()
            if mark_executed("BOOKMARK", name):
                self.open_bookmark(name)
                need_api_call = True

        # --- ВИКИПЕДИЯ ---
        for query in re.findall(r'\[WIKI\](.*?)\[/WIKI\]', text, re.DOTALL):
            query = query.strip()
            if mark_executed("WIKI", query):
                self.process_search(query, "wiki")
                need_api_call = True

        # --- НОВОСТИ ---
        for query in re.findall(r'\[NEWS\](.*?)\[/NEWS\]', text, re.DOTALL):
            query = query.strip()
            if mark_executed("NEWS", query):
                self.process_search(query, "news")
                need_api_call = True

        # --- ОБЩИЙ ПОИСК ---
        for query in re.findall(r'\[SEARCH\](.*?)\[/SEARCH\]', text, re.DOTALL):
            query = query.strip()
            if mark_executed("SEARCH", query):
                self.process_search(query, "general")
                need_api_call = True

        # --- ПОИСК В БРАУЗЕРЕ (ЗАПАСНОЙ) ---
        for query in re.findall(r'\[BROWSER_SEARCH\](.*?)\[/BROWSER_SEARCH\]', text, re.DOTALL):
            query = query.strip()
            if mark_executed("BROWSER_SEARCH", query):
                self.google_search_and_extract(query)
                need_api_call = True

        # --- ОТКРЫТИЕ ПРИЛОЖЕНИЯ ---
        for app in re.findall(r'\[OPEN_APP\](.*?)\[/OPEN_APP\]', text, re.DOTALL):
            app = app.strip()
            if not mark_executed("OPEN_APP", app):
                continue
            need_api_call = True
            if self.android_bridge:
                try:
                    self.android_bridge.openApp(app)
                    self.android_bridge.playOpen()
                    self.messages.append({"role": "system", "content": f"✅ Приложение '{app}' открыто."})
                except Exception as e:
                    self.messages.append({"role": "system", "content": f"❌ Не удалось открыть приложение '{app}': {e}"})
                    if self.android_bridge:
                        self.android_bridge.playError()
            else:
                self.messages.append({"role": "system", "content": "❌ Нет доступа к Android Bridge."})

        # --- СОЗДАНИЕ ФАЙЛА ---
        for match in re.findall(r'\[NOTEPAD\](.*?)\[/NOTEPAD\]', text, re.DOTALL):
            lines = match.strip().split('\n', 1)
            filename = lines[0].strip()
            content = lines[1] if len(lines) > 1 else ""
            if not mark_executed("NOTEPAD", filename):
                continue
            need_api_call = True
            actual_path = os.path.expandvars(filename)
            try:
                os.makedirs(os.path.dirname(actual_path), exist_ok=True)
                with open(actual_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                self.messages.append({"role": "system", "content": f"✅ Файл '{actual_path}' создан."})
                if self.android_bridge:
                    self.android_bridge.playDone()
            except Exception as e:
                self.messages.append({"role": "system", "content": f"❌ Ошибка создания файла '{actual_path}': {e}"})
                if self.android_bridge:
                    self.android_bridge.playError()

        # --- ДОБАВЛЕНИЕ В КОНЕЦ ФАЙЛА ---
        for match in re.findall(r'\[NOTEPAD_APPEND\](.*?)\[/NOTEPAD_APPEND\]', text, re.DOTALL):
            lines = match.strip().split('\n', 1)
            filename = lines[0].strip()
            content = lines[1] if len(lines) > 1 else ""
            if not mark_executed("NOTEPAD_APPEND", filename):
                continue
            need_api_call = True
            actual_path = os.path.expandvars(filename)
            try:
                os.makedirs(os.path.dirname(actual_path), exist_ok=True)
                with open(actual_path, 'a', encoding='utf-8') as f:
                    f.write(content)
                self.messages.append({"role": "system", "content": f"✅ Содержимое добавлено в конец файла '{actual_path}'."})
                if self.android_bridge:
                    self.android_bridge.playDone()
            except Exception as e:
                self.messages.append({"role": "system", "content": f"❌ Ошибка добавления в файл '{actual_path}': {e}"})
                if self.android_bridge:
                    self.android_bridge.playError()

        # --- ДОБАВЛЕНИЕ В НАЧАЛО ФАЙЛА ---
        for match in re.findall(r'\[NOTEPAD_PREPEND\](.*?)\[/NOTEPAD_PREPEND\]', text, re.DOTALL):
            lines = match.strip().split('\n', 1)
            filename = lines[0].strip()
            content = lines[1] if len(lines) > 1 else ""
            if not mark_executed("NOTEPAD_PREPEND", filename):
                continue
            need_api_call = True
            actual_path = os.path.expandvars(filename)
            try:
                os.makedirs(os.path.dirname(actual_path), exist_ok=True)
                existing = ""
                if os.path.exists(actual_path):
                    with open(actual_path, 'r', encoding='utf-8') as f:
                        existing = f.read()
                with open(actual_path, 'w', encoding='utf-8') as f:
                    f.write(content + existing)
                self.messages.append({"role": "system", "content": f"✅ Содержимое добавлено в начало файла '{actual_path}'."})
                if self.android_bridge:
                    self.android_bridge.playDone()
            except Exception as e:
                self.messages.append({"role": "system", "content": f"❌ Ошибка добавления в начало файла '{actual_path}': {e}"})
                if self.android_bridge:
                    self.android_bridge.playError()

        # --- ЗАМЕНА ТЕКСТА ---
        for match in re.findall(r'\[NOTEPAD_REPLACE\](.*?)\[/NOTEPAD_REPLACE\]', text, re.DOTALL):
            parts = match.strip().split('\n', 2)
            if len(parts) < 3:
                continue
            filename = parts[0].strip()
            old = parts[1]
            new = parts[2]
            key = f"{filename}|{old}|{new}"
            if not mark_executed("NOTEPAD_REPLACE", key):
                continue
            need_api_call = True
            actual_path = os.path.expandvars(filename)
            try:
                if os.path.exists(actual_path):
                    with open(actual_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    new_content = content.replace(old, new)
                    if new_content != content:
                        with open(actual_path, 'w', encoding='utf-8') as f:
                            f.write(new_content)
                        self.messages.append({"role": "system", "content": f"✅ Текст в файле '{actual_path}' заменён (старое: '{old}', новое: '{new}')."})
                        if self.android_bridge:
                            self.android_bridge.playDone()
                    else:
                        self.messages.append({"role": "system", "content": f"⚠️ Текст '{old}' не найден в файле '{actual_path}'."})
                        if self.android_bridge:
                            self.android_bridge.playError()
                else:
                    self.messages.append({"role": "system", "content": f"❌ Файл '{actual_path}' не существует."})
                    if self.android_bridge:
                        self.android_bridge.playError()
            except Exception as e:
                self.messages.append({"role": "system", "content": f"❌ Ошибка замены текста в '{actual_path}': {e}"})
                if self.android_bridge:
                    self.android_bridge.playError()

        # --- УДАЛЕНИЕ СТРОКИ ---
        for match in re.findall(r'\[NOTEPAD_DELETE_LINE\](.*?)\[/NOTEPAD_DELETE_LINE\]', text, re.DOTALL):
            parts = match.strip().split('\n', 1)
            if len(parts) < 2:
                continue
            filename = parts[0].strip()
            try:
                line_num = int(parts[1].strip())
            except ValueError:
                continue
            key = f"{filename}|{line_num}"
            if not mark_executed("NOTEPAD_DELETE_LINE", key):
                continue
            need_api_call = True
            actual_path = os.path.expandvars(filename)
            try:
                if os.path.exists(actual_path):
                    with open(actual_path, 'r', encoding='utf-8') as f:
                        lines = f.readlines()
                    if 1 <= line_num <= len(lines):
                        del lines[line_num - 1]
                        with open(actual_path, 'w', encoding='utf-8') as f:
                            f.writelines(lines)
                        self.messages.append({"role": "system", "content": f"✅ Строка {line_num} удалена из файла '{actual_path}'."})
                        if self.android_bridge:
                            self.android_bridge.playDone()
                    else:
                        self.messages.append({"role": "system", "content": f"⚠️ Строка {line_num} не существует в файле '{actual_path}'."})
                        if self.android_bridge:
                            self.android_bridge.playError()
                else:
                    self.messages.append({"role": "system", "content": f"❌ Файл '{actual_path}' не существует."})
                    if self.android_bridge:
                        self.android_bridge.playError()
            except Exception as e:
                self.messages.append({"role": "system", "content": f"❌ Ошибка удаления строки из '{actual_path}': {e}"})
                if self.android_bridge:
                    self.android_bridge.playError()

        # --- БУФЕР ОБМЕНА: УСТАНОВКА ---
        for match in re.findall(r'\[CLIPBOARD_SET\](.*?)\[/CLIPBOARD_SET\]', text, re.DOTALL):
            txt = match.strip()
            if not mark_executed("CLIPBOARD_SET", txt):
                continue
            need_api_call = True
            if self.android_bridge:
                try:
                    self.android_bridge.setClipboard(txt)
                    self.messages.append({"role": "system", "content": f"✅ Текст скопирован в буфер обмена: '{txt}'"})
                    if self.android_bridge:
                        self.android_bridge.playDone()
                except Exception as e:
                    self.messages.append({"role": "system", "content": f"❌ Ошибка копирования в буфер: {e}"})
                    if self.android_bridge:
                        self.android_bridge.playError()
            else:
                self.messages.append({"role": "system", "content": "❌ Нет доступа к Android Bridge."})

        # --- БУФЕР ОБМЕНА: ПОЛУЧЕНИЕ ---
        if '[CLIPBOARD_GET]' in text and mark_executed("CLIPBOARD_GET", ""):
            need_api_call = True
            if self.android_bridge:
                try:
                    content = self.android_bridge.getClipboard()
                    if content:
                        self.messages.append({"role": "system", "content": f"📋 Содержимое буфера обмена: {content}"})
                    else:
                        self.messages.append({"role": "system", "content": "📋 Буфер обмена пуст."})
                    if self.android_bridge:
                        self.android_bridge.playDone()
                except Exception as e:
                    self.messages.append({"role": "system", "content": f"❌ Ошибка чтения буфера: {e}"})
                    if self.android_bridge:
                        self.android_bridge.playError()
            else:
                self.messages.append({"role": "system", "content": "❌ Нет доступа к Android Bridge."})

        # --- КАЛЬКУЛЯТОР ---
        for expr in re.findall(r'\[CALC\](.*?)\[/CALC\]', text, re.DOTALL):
            expr = expr.strip()
            if not mark_executed("CALC", expr):
                continue
            need_api_call = True
            try:
                allowed = {"abs": abs, "round": round, "min": min, "max": max}
                code = compile(expr, "<string>", "eval")
                for name in code.co_names:
                    if name not in allowed:
                        raise NameError(f"Использование {name} запрещено")
                result = eval(expr, {"__builtins__": {}}, allowed)
                self.messages.append({"role": "system", "content": f"🧮 Результат вычисления '{expr}' = {result}"})
                if self.android_bridge:
                    self.android_bridge.playDone()
            except Exception as e:
                self.messages.append({"role": "system", "content": f"❌ Ошибка вычисления '{expr}': {e}"})
                if self.android_bridge:
                    self.android_bridge.playError()

        # --- ОТКРЫТИЕ ССЫЛКИ ---
        for url in re.findall(r'\[OPEN_URL\](.*?)\[/OPEN_URL\]', text, re.DOTALL):
            url = url.strip()
            if not mark_executed("OPEN_URL", url):
                continue
            need_api_call = True
            if self.android_bridge:
                try:
                    self.android_bridge.openUrl(url)
                    self.android_bridge.playOpen()
                    self.messages.append({"role": "system", "content": f"🔗 Ссылка открыта: {url}"})
                    if self.android_bridge:
                        self.android_bridge.playDone()
                except Exception as e:
                    self.messages.append({"role": "system", "content": f"❌ Не удалось открыть ссылку {url}: {e}"})
                    if self.android_bridge:
                        self.android_bridge.playError()
            else:
                self.messages.append({"role": "system", "content": "❌ Нет доступа к Android Bridge."})

        # --- НАПОМИНАНИЕ ---
        for remind in re.findall(r'\[REMIND\](.*?)\[/REMIND\]', text, re.DOTALL):
            parts = remind.strip().split(';', 1)
            if len(parts) != 2:
                continue
            sec_str, msg = parts[0].strip(), parts[1].strip()
            key = f"{sec_str}|{msg}"
            if not mark_executed("REMIND", key):
                continue
            need_api_call = True
            try:
                sec = int(sec_str)
                threading.Thread(target=self._reminder, args=(sec, msg), daemon=True).start()
                self.messages.append({"role": "system", "content": f"⏰ Напоминание установлено через {sec} секунд: '{msg}'"})
                if self.android_bridge:
                    self.android_bridge.playDone()
            except ValueError:
                self.messages.append({"role": "system", "content": "❌ Неверный формат напоминания. Используйте: секунды;сообщение"})
                if self.android_bridge:
                    self.android_bridge.playError()

        # --- ГРОМКОСТЬ ---
        for vol in re.findall(r'\[VOLUME\](\d+)\[/VOLUME\]', text, re.DOTALL):
            if not mark_executed("VOLUME", vol):
                continue
            need_api_call = True
            if self.android_bridge:
                try:
                    self.android_bridge.setVolume(int(vol))
                    self.messages.append({"role": "system", "content": f"🔊 Громкость установлена на {vol}%"})
                    if self.android_bridge:
                        self.android_bridge.playDone()
                except Exception as e:
                    self.messages.append({"role": "system", "content": f"❌ Ошибка установки громкости: {e}"})
                    if self.android_bridge:
                        self.android_bridge.playError()
            else:
                self.messages.append({"role": "system", "content": "❌ Нет доступа к Android Bridge."})

        # --- ЧТЕНИЕ ФАЙЛА ---
        for filepath in re.findall(r'\[READ_FILE\](.*?)\[/READ_FILE\]', text, re.DOTALL):
            filepath = filepath.strip()
            if mark_executed("READ_FILE", filepath):
                self.read_file_from_disk(filepath)  # внутри добавляет сообщение и сохраняет историю
                need_api_call = True

        # --- ЗАПИСЬ ФАЙЛА ---
        for match in re.findall(r'\[WRITE_FILE\](.*?)\[/WRITE_FILE\]', text, re.DOTALL):
            parts = match.strip().split('\n', 1)
            if len(parts) != 2:
                need_api_call = True
                self.messages.append({"role": "system", "content": "❌ Неверный формат [WRITE_FILE]. Используйте: путь\nсодержимое"})
                if self.android_bridge:
                    self.android_bridge.playError()
                continue
            filepath = parts[0].strip()
            content = parts[1]
            if mark_executed("WRITE_FILE", filepath):
                self.write_file_to_disk(filepath, content)  # внутри добавляет сообщение
                need_api_call = True

        # --- ПЕРЕВОД ---
        for lang_code, original in re.findall(r'\[TRANSLATE(?:\s+([a-z]{2}(?:-[a-z]{2})?))?\](.*?)\[/TRANSLATE\]', text, re.DOTALL):
            original = original.strip()
            target_lang = lang_code if lang_code else 'en'
            key = f"{target_lang}|{original}"
            if mark_executed("TRANSLATE", key):
                need_api_call = True
                threading.Thread(target=self.translate_text, args=(original, target_lang), daemon=True).start()
                self.messages.append({"role": "system", "content": f"🌐 Запущен перевод на {target_lang}..."})

        # --- ГЕНЕРАЦИЯ ИЗОБРАЖЕНИЙ ---
        for match in re.findall(r'\[GENERATE_IMAGE\](.*?)\[/GENERATE_IMAGE\]', text, re.DOTALL):
            prompt = match.strip()
            if mark_executed("GENERATE_IMAGE", prompt):
                need_api_call = True
                threading.Thread(target=self.generate_image, args=(prompt,), daemon=True).start()
                self.messages.append({"role": "system", "content": f"🎨 Запущена генерация изображения по запросу: {prompt}"})
                self.save_history()
        
        # --- МУЛЬТИМЕДИА — ЗАЩИТА ОТ ДУБЛЕЙ ---
        now = time.time()
        if now - self._last_media_time < 5:
            pass  # Пропускаем все медиа-команды если прошло <5 секунд
        else:
            self._last_media_time = now
            if any(tag in text for tag in ['[MEDIA_NEXT]']):
                if mark_executed("MEDIA_NEXT", ""):
                    need_api_call = True
                    if self.android_bridge:
                        try:
                            self.android_bridge.nextTrack()
                            self.android_bridge.playDone()
                            self.messages.append({"role": "system", "content": "⏭️ Следующий трек"})
                        except Exception as e:
                            self.messages.append({"role": "system", "content": f"❌ Ошибка переключения трека: {e}"})
                            if self.android_bridge:
                                self.android_bridge.playError()
                    else:
                        self.messages.append({"role": "system", "content": "❌ Нет доступа к Android Bridge."})
            
            if any(tag in text for tag in ['[MEDIA_PREV]']):
                if mark_executed("MEDIA_PREV", ""):
                    need_api_call = True
                    if self.android_bridge:
                        try:
                            self.android_bridge.prevTrack()
                            self.android_bridge.playDone()
                            self.messages.append({"role": "system", "content": "⏮️ Предыдущий трек"})
                        except Exception as e:
                            self.messages.append({"role": "system", "content": f"❌ Ошибка переключения трека: {e}"})
                            if self.android_bridge:
                                self.android_bridge.playError()
                    else:
                        self.messages.append({"role": "system", "content": "❌ Нет доступа к Android Bridge."})

            if any(tag in text for tag in ['[MEDIA_PLAY_PAUSE]']):
                if mark_executed("MEDIA_PLAY_PAUSE", ""):
                    need_api_call = True
                    if self.android_bridge:
                        try:
                            self.android_bridge.playPause()
                            self.android_bridge.playDone()
                            self.messages.append({"role": "system", "content": "⏯️ Пауза/Воспроизведение"})
                        except Exception as e:
                            self.messages.append({"role": "system", "content": f"❌ Ошибка паузы/воспроизведения: {e}"})
                            if self.android_bridge:
                                self.android_bridge.playError()
                    else:
                        self.messages.append({"role": "system", "content": "❌ Нет доступа к Android Bridge."})

        # --- ФИНАЛЬНЫЙ ВЫЗОВ API ТОЛЬКО ЕСЛИ БЫЛИ КОМАНДЫ ---
        if need_api_call:
            self.save_history()

    def translate_text(self, text, target_lang='en'):
        try:
            translated = GoogleTranslator(source='auto', target=target_lang).translate(text)
            self.messages.append({"role": "system", "content": f"Перевод ({target_lang}): {translated}"})
            self.save_history()
            threading.Thread(target=self._call_api, daemon=True).start()
        except Exception as e:
            self.messages.append({"role": "system", "content": f"Ошибка перевода: {e}"})
            self.save_history()
            threading.Thread(target=self._call_api, daemon=True).start()

    def _reminder(self, seconds, message):
        time.sleep(seconds)
        if self.android_bridge:
            self.android_bridge.addAssistantMessage(f"⏰ НАПОМИНАНИЕ: {message}")
            self.messages.append({"role": "system", "content": f"⏰ НАПОМИНАНИЕ: {message}"})
            self.save_history()
            threading.Thread(target=self._call_api, daemon=True).start()

    def process_message(self, user_text, activity):
        if user_text.strip() == '/clearhistory':
            if self.messages and self.messages[0]["role"] == "system":
                self.messages = [self.messages[0]]
            else:
                self.messages = []
            self.save_history()
            if self.android_bridge:
                self.android_bridge.addAssistantMessage("История очищена.")
            return

        self.messages.append({"role": "user", "content": user_text})
        self.save_history()
        
        lower_input = user_text.lower()
        
        # Проверяем, что сообщение состоит ТОЛЬКО из фразы "привет дорогая" (с возможной запятой)
        if re.fullmatch(r'привет\s*,?\s*дорогая', lower_input.strip()):
            if self.android_bridge and not self._greeting_played:
                self.android_bridge.playGreeting()  
                self._greeting_played = True
        
        # Создать заметку в приложении MIUI Notes (без разделителя)
        if re.search(r'создай заметку\s+(.+)', user_text, re.IGNORECASE):
            text = re.search(r'создай заметку\s+(.+)', user_text, re.IGNORECASE).group(1)
            result = self.create_note_in_miui(text)
            self.messages.append({"role": "system", "content": result})
            self.save_history()
            return

        # Создать текстовый файл в папке Загрузки
        if re.search(r'создай заметку в загрузках\s+(.+)', user_text, re.IGNORECASE):
            text = re.search(r'создай заметку в загрузках\s+(.+)', user_text, re.IGNORECASE).group(1)
            filename = f"заметка_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
            result = self.create_text_file_in_downloads(filename, text)
            self.messages.append({"role": "system", "content": result})
            self.save_history()
            return

        # Включить трек
        if re.search(r'включи трек\s+(.+)', user_text, re.IGNORECASE):
            track = re.search(r'включи трек\s+(.+)', user_text, re.IGNORECASE).group(1)
            result = self.play_track_in_miui(track)
            self.messages.append({"role": "system", "content": result})
            self.save_history()
            return

        # Открыть браузер и найти (запасной вариант)
        if re.search(r'открой браузер и найди\s+(.+)', user_text, re.IGNORECASE):
            query = re.search(r'открой браузер и найди\s+(.+)', user_text, re.IGNORECASE).group(1)
            self.google_search_and_extract(query)
            return

        # Создать событие в календаре
        if re.search(r'создай событие\s+(.+)', user_text, re.IGNORECASE):
            title = re.search(r'создай событие\s+(.+)', user_text, re.IGNORECASE).group(1)
            # Пока дата не передаётся, по умолчанию будет сегодняшняя
            result = self.create_calendar_event(title, "сегодня")
            self.messages.append({"role": "system", "content": result})
            self.save_history()
            return
        
        # Поиск в интернете через SearXNG (быстрый) - только если сказано "в интернете" или "в сети"
        search_internet_match = re.search(r'(?:найди|поищи|ищу)\s+в\s+(?:интернете|сети)\s+(.+)', user_text, re.IGNORECASE)
        if search_internet_match:
            query = search_internet_match.group(1).strip()
            if self.android_bridge:
                self.android_bridge.playSearch()
            if query:
                self.process_search(query, "general")
            else:
                self.messages.append({"role": "system", "content": "Не поняла, что искать в интернете."})
                self.save_history()
                self._call_api()
            return
        
        # Открыть приложение — СНАЧАЛА MIUI, ПОТОМ ПОИСК
        if any(kw in lower_input for kw in ['открой', 'открыть', 'запусти']):
            app_match = re.search(r'(?:открой|открыть|запусти)\s+(.+)', lower_input)
            if app_match:
                app_name = app_match.group(1).strip().lower()
                package = None
                
                # 1. Ищем в словаре MIUI
                for key, pkg in self.miui_apps.items():
                    if key in app_name or app_name in key:
                        package = pkg
                        break
                
                # 2. Если не нашли — ищем через Android Bridge
                if not package and self.android_bridge:
                    package = self.android_bridge.findAppPackage(app_name)
                
                # 3. Если всё ещё нет — используем как есть
                if not package:
                    package = app_name
                
                tag = f"[OPEN_APP]{package}[/OPEN_APP]"
                self.messages.append({"role": "assistant", "content": tag})
                self.save_history()
                self._process_commands(tag)
            return
        
        # Найти файл
        if any(kw in lower_input for kw in ['найди файл', 'поищи файл', 'найди документ']):
            file_match = re.search(r'(?:найди|поищи)\s+файл\s+(.+)', lower_input)
            if file_match:
                filename = file_match.group(1).strip()
                self.find_file_in_storage(filename)
            else:
                self.messages.append({"role": "system", "content": "Не указано имя файла."})
                self.save_history()
                self._call_api()
            return

        # Открыть файл (по пути или по имени)
        if any(kw in lower_input for kw in ['открой файл', 'открой документ', 'открой фото', 'открой картинку']):
            file_match = re.search(r'(?:открой|покажи)\s+(?:файл|документ|фото|картинку)\s+(.+)', lower_input)
            if file_match:
                file_ref = file_match.group(1).strip()
                if '/' in file_ref or file_ref.startswith('storage'):
                    if self.android_bridge:
                        tag = f"[OPEN_URL]file://{file_ref}[/OPEN_URL]"
                        self.messages.append({"role": "assistant", "content": tag})
                        self.save_history()
                        self._process_commands(tag)
                else:
                    self.open_file_by_name(file_ref)
            else:
                self.messages.append({"role": "system", "content": "Не указан файл."})
                self.save_history()
                self._call_api()
            return

        # Открыть ссылку
        url_match = re.search(r'(?:открой|перейди на)\s+(https?://\S+)', user_text)
        if url_match:
            url = url_match.group(1)
            tag = f"[OPEN_URL]{url}[/OPEN_URL]"
            self.messages.append({"role": "assistant", "content": tag})
            self.save_history()
            self._process_commands(tag)
            return

        # Записать файл
        if any(kw in lower_input for kw in ['запиши', 'сохрани', 'создай файл']):
            file_match = re.search(r'(?:запиши|сохрани|создай файл)\s+([^\s]+)\s+(.+)', user_text)
            if file_match:
                filename = file_match.group(1).strip()
                content = file_match.group(2).strip()
                if '/' in filename or '\\' in filename:
                    filepath = filename
                else:
                    if self.android_bridge:
                        downloads = self.android_bridge.getDownloadsPath()
                        filepath = os.path.join(downloads, filename)
                    else:
                        filepath = filename
                tag = f"[NOTEPAD]{filepath}\n{content}[/NOTEPAD]"
                self.messages.append({"role": "assistant", "content": tag})
                self.save_history()
                self._process_commands(tag)
            else:
                self.messages.append({"role": "system", "content": "Неверный формат. Используй: запиши имя_файла содержимое"})
                self.save_history()
                self._call_api()
            return

        # Прочитать файл
        if any(kw in lower_input for kw in ['прочитай', 'покажи файл', 'открой файл']):
            file_match = re.search(r'(?:прочитай|покажи|открой)\s+(?:файл\s+)?(\S+)', user_text)
            if file_match:
                filepath = file_match.group(1)
                tag = f"[READ_FILE]{filepath}[/READ_FILE]"
                self.messages.append({"role": "assistant", "content": tag})
                self.save_history()
                self._process_commands(tag)
            else:
                self.messages.append({"role": "system", "content": "Не указан файл."})
                self.save_history()
                self._call_api()
            return

        # Генерация изображения по текстовому запросу
        if re.search(r'(сгенерируй|нарисуй|создай изображение)\s+(.+)', user_text, re.IGNORECASE):
            prompt = re.search(r'(сгенерируй|нарисуй|создай изображение)\s+(.+)', user_text, re.IGNORECASE).group(2)
            tag = f"[GENERATE_IMAGE]{prompt}[/GENERATE_IMAGE]"
            self.messages.append({"role": "assistant", "content": tag})
            self.save_history()
            self._process_commands(tag)
            return

        # Воспроизведение медиафайла
        if any(kw in user_text.lower() for kw in ['включи', 'запусти', 'воспроизведи', 'сыграй']):
            file_match = re.search(r'(?:включи|запусти|воспроизведи|сыграй)\s+(.+?)(?:\s|$)', user_text.lower())
            if file_match:
                filename = file_match.group(1).strip()
                if self.android_bridge:
                    results = self.android_bridge.searchFiles(filename)
                    if results:
                        file_to_play = results[0]
                        if isinstance(file_to_play, (list, tuple)):
                            file_to_play = file_to_play[0]
                        self._play_media_file(file_to_play)
                        self.messages.append({"role": "system", "content": f"🎵 Воспроизвожу {filename}"})
                    else:
                        self.messages.append({"role": "system", "content": f"❌ Файл {filename} не найден."})
                    self.save_history()
                    self._call_api()
                    return

        self._call_api()

galya_instance = None

def process_image(base64_image, description=""):
        global galya_instance
        if galya_instance is None:
            galya_instance = Galya(API_KEY)
        galya_instance.process_image(base64_image, description)

def process_uploaded_text(filename, content):
    global galya_instance
    if galya_instance is None:
        galya_instance = Galya(API_KEY)
    galya_instance.process_uploaded_text(filename, content)

def set_bridge(bridge):
        global galya_instance
        if galya_instance is None:
            galya_instance = Galya(API_KEY)
        galya_instance.android_bridge = bridge

def process_message(user_text, activity):
        global galya_instance
        if galya_instance is None:
            galya_instance = Galya(API_KEY)
        galya_instance.process_message(user_text, activity)

def get_conversation_history():
    global galya_instance
    if galya_instance is None:
        return []
    history = []
    for msg in galya_instance.messages:
        role = msg.get("role")
        content = msg.get("content")
        # Показываем только сообщения пользователя и ассистента (не системные)
        if role in ("user", "assistant") and isinstance(content, str):
            history.append({"role": role, "content": content})
    return history